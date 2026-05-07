"""
员工离职数据分析系统 - Streamlit Cloud版本 V3.2
"""
import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import re
from io import BytesIO
import json
import requests
import os
from datetime import datetime

import warnings
warnings.filterwarnings("ignore")

# DeepSeek API 配置
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "sk-3b6b0ab65ce148a2a9f48abcf007cb5b")
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

def get_download_config(name):
    return {
        "toImageButtonOptions": {
            "format": "png",
            "filename": name,
            "height": 600,
            "width": 1200,
            "scale": 2
        },
        "displaylogo": False,
        "modeBarButtonsToRemove": ["select2d", "lasso2d"]
    }

st.set_page_config(page_title="员工离职数据分析", layout="wide")

COLORS = {"primary": "#4F46E5", "danger": "#EF4444", "warning": "#F59E0B", "purple": "#8B5CF6", "success": "#10B981"}

def get_month_str(n):
    return "2026年" + str(n).zfill(2) + "月"

def cat_tenure(y):
    if pd.isna(y): return "未知"
    elif y <= 0.5: return "0.5年及以下"
    elif y <= 1: return "0.5-1年"
    elif y <= 3: return "1-3年"
    elif y <= 5: return "3-5年"
    else: return "5年以上"

def cat_level(level):
    if pd.isna(level): return "未知"
    level = str(level).strip()
    if level.startswith("中"): return "中级"
    elif level.startswith("高"): return "高级"
    elif level.startswith("初"): return "初级"
    elif level.startswith("资深"): return "资深"
    elif level.startswith("总监"): return "总监"
    else: return level

def call_deepseek(prompt, system_prompt=None):
    """调用 DeepSeek API"""
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}"
    }
    
    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": prompt})
    
    payload = {
        "model": "deepseek-chat",
        "messages": messages,
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=60)
        if response.status_code == 200:
            result = response.json()
            return result["choices"][0]["message"]["content"]
        else:
            return f"API调用失败: {response.status_code} - {response.text}"
    except Exception as e:
        return f"API调用错误: {str(e)}"

def export_to_pdf(analysis_text, month):
    """将分析报告导出为 PDF"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Title_CN', fontName='Helvetica', fontSize=18, alignment=1, spaceAfter=20))
    styles.add(ParagraphStyle(name='Heading1_CN', fontName='Helvetica-Bold', fontSize=14, spaceBefore=15, spaceAfter=8))
    styles.add(ParagraphStyle(name='Body_CN', fontName='Helvetica', fontSize=10, leading=14, spaceAfter=6))
    
    story = []
    story.append(Paragraph(f"员工离职分析报告", styles['Title_CN']))
    story.append(Paragraph(f"{month} | 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Body_CN']))
    story.append(Spacer(1, 0.5*cm))
    
    lines = analysis_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.2*cm))
            continue
        if line.startswith('## '):
            story.append(Paragraph(line.replace('## ', ''), styles['Heading1_CN']))
        elif line.startswith('# '):
            story.append(Paragraph(line, styles['Title_CN']))
        elif line.startswith('**') and line.endswith('**'):
            story.append(Paragraph(line, styles['Heading1_CN']))
        elif line.startswith('- ') or line.startswith('* '):
            story.append(Paragraph(line, styles['Body_CN']))
        elif line.startswith('>'):
            story.append(Paragraph(line.replace('>', '').strip(), styles['Body_CN']))
        elif line.startswith('|'):
            continue
        else:
            story.append(Paragraph(line, styles['Body_CN']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def export_to_word(analysis_text, month, data_summary=None):
    """将分析报告导出为 Word"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    title = doc.add_heading('员工离职分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{month} | 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    lines = analysis_text.split('\n')
    in_table = False
    table_rows = []
    
    for line in lines:
        line = line.strip()
        if not line:
            if in_table and table_rows:
                in_table = False
            continue
        
        if line.startswith('|') and line.endswith('|'):
            if '---' not in line:
                in_table = True
                cells = [c.strip() for c in line.split('|')[1:-1]]
                table_rows.append(cells)
            continue
        else:
            if in_table and table_rows:
                if table_rows:
                    tbl = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    tbl.style = 'Table Grid'
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for i, row in enumerate(table_rows):
                        for j, cell_text in enumerate(row):
                            cell = tbl.rows[i].cells[j]
                            cell.text = cell_text
                            for para in cell.paragraphs:
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                table_rows = []
                in_table = False
        
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.bold = True
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line, style='List Bullet')
        elif line.startswith('>'):
            doc.add_paragraph(line.replace('>', '').strip())
        elif in_table:
            pass
        else:
            doc.add_paragraph(line)
    
    if data_summary:
        doc.add_page_break()
        doc.add_heading('原始数据摘要', level=1)
        p = doc.add_paragraph()
        run = p.add_run(data_summary)
        run.font.size = Pt(10)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def get_data_summary(r):
    """获取数据摘要文本"""
    lines = []
    lines.append(f"分析周期: {r['month']}")
    lines.append(f"离职率: {r['rate']}%")
    lines.append(f"离职总人数: {r['turn']} 人")
    lines.append(f"平均人数: {r['avg']} 人")
    if r["type"]:
        for t in r["type"]:
            lines.append(f"{t.get('离职类型', '类型')}: {t.get('人数', 0)} 人")
    if r["dept"]:
        lines.append("\n各部门离职数据:")
        for d in r["dept"][:10]:
            lines.append(f"- {d['一级组织']}: 离职{d['离职人数']}人, 离职率{d['离职率']}%")
    if r["reason"]:
        lines.append("\n离职原因分布:")
        for reason in r["reason"][:8]:
            lines.append(f"- {reason['离职原因']}: {reason['人数']}人, 占比{reason['占比']}%")
    return '\n'.join(lines)

def generate_ai_analysis(r):
    """使用 DeepSeek AI 生成智能分析报告"""
    # 准备数据摘要
    主动离职 = 0
    被动离职 = 0
    if r["type"]:
        for t in r["type"]:
            if "主动" in str(t.get("离职类型", "")):
                主动离职 = t.get("人数", 0)
            elif "被动" in str(t.get("离职类型", "")):
                被动离职 = t.get("人数", 0)
    
    # 构建数据摘要文本
    data_summary = f"""
分析周期: {r['month']}
离职率: {r['rate']}%
离职总人数: {r['turn']} 人
平均人数: {r['avg']} 人
主动离职: {主动离职} 人
被动离职: {被动离职} 人
"""
    
    # 部门数据
    if r["dept"]:
        data_summary += "\n各部门离职数据:\n"
        for d in r["dept"]:
            data_summary += f"- {d['一级组织']}: 离职{d['离职人数']}人, 离职率{d['离职率']}%\n"
    
    # 离职原因
    if r["reason"]:
        data_summary += "\n离职原因分布:\n"
        for reason in r["reason"][:8]:
            data_summary += f"- {reason['离职原因']}: {reason['人数']}人, 占比{reason['占比']}%\n"
    
    # 司龄分布
    if r["tenure"]:
        data_summary += "\n司龄分布:\n"
        for t in r["tenure"]:
            data_summary += f"- {t['司龄段']}: {t['人数']}人, 占比{t['占比']}%\n"
    
    # 职级分布
    if r["level"]:
        data_summary += "\n职级分布:\n"
        for l in r["level"]:
            data_summary += f"- {l['职级合并']}: {l['人数']}人, 占比{l['占比']}%\n"
    
    # DeepSeek 分析提示词
    system_prompt = """你是一位资深的人力资源数据分析专家，擅长分析员工离职数据并提供专业、可执行的管理建议。
请用 Markdown 格式输出分析报告，包含以下维度：
1. 整体概况（简洁摘要）
2. 主动离职分析（原因+建议）
3. 被动离职分析（原因+建议）
4. 部门风险分析（高风险部门识别）
5. 司龄分布洞察
6. 职级分布分析（中坚力量关注）
7. 风险预警（自动识别高风险信号）
8. 综合改善建议（表格形式，含优先级）

注意：
- 使用中文输出
- 建议要具体、可执行
- 风险预警要客观量化
- 控制在1500字以内
- 用 emoji 增强可读性"""

    user_prompt = f"请分析以下离职数据并生成报告：\n{data_summary}"
    
    # 调用 DeepSeek API
    return call_deepseek(user_prompt, system_prompt)

class Proc:
    def __init__(self, xlsx):
        self.period = {}
        self.turn = None
        for sh in xlsx.sheet_names:
            try:
                df = pd.read_excel(xlsx, sheet_name=sh)
                if "离职" in sh and "期初" not in sh and "期末" not in sh:
                    self.turn = df.copy()
                    m = re.search(r"^(\d{1,2})月", sh)
                    if m:
                        self.turn["离职月份"] = get_month_str(int(m.group(1)))
                    elif "最后工作日" in df.columns:
                        self.turn["最后工作日"] = pd.to_datetime(self.turn["最后工作日"], errors="coerce")
                        self.turn["离职月份"] = self.turn["最后工作日"].dt.strftime("%Y年%m月")
                    continue
                m = re.search(r"^(\d{1,2})月", sh)
                if not m: continue
                om = int(m.group(1))
                inner = re.search(r"（(\d{1,2})月[^）]*期末", sh)
                if inner:
                    pm = int(inner.group(1))
                    if pm not in self.period: self.period[pm] = {"s": None, "e": None}
                    self.period[pm]["e"] = df
                elif "期末" in sh:
                    if om not in self.period: self.period[om] = {"s": None, "e": None}
                    self.period[om]["e"] = df
                if "期初" in sh:
                    if om not in self.period: self.period[om] = {"s": None, "e": None}
                    self.period[om]["s"] = df
            except: pass
    
    def months(self):
        return sorted([m for m in self.period.keys() if self.period[m]["s"] is not None])
    
    def analyze(self, mns):
        if not isinstance(mns, list): mns = [mns]
        ts, te, tt = 0, 0, 0
        all_t = pd.DataFrame()
        ds, de, dt = {}, {}, {}
        for mn in mns:
            s, e = self.period.get(mn, {}).get("s"), self.period.get(mn, {}).get("e")
            ms = get_month_str(mn)
            if s is not None:
                ts += len(s)
                if "一级组织" in s.columns:
                    for d, c in s.groupby("一级组织").size().items(): ds[d] = ds.get(d, 0) + c
            if e is not None and len(e) > 0:
                te += len(e)
                if "一级组织" in e.columns:
                    for d, c in e.groupby("一级组织").size().items(): de[d] = de.get(d, 0) + c
            if self.turn is not None and "离职月份" in self.turn.columns:
                mt = self.turn[self.turn["离职月份"] == ms].copy()
                tt += len(mt)
                all_t = pd.concat([all_t, mt], ignore_index=True)
                if "一级组织" in mt.columns:
                    for d, c in mt.groupby("一级组织").size().items(): dt[d] = dt.get(d, 0) + c
        avg = (ts + te) / 2 if (ts + te) > 0 else 0
        rate = round((tt / avg * 100), 2) if avg > 0 else 0
        lbl = get_month_str(mns[0]) if len(mns) == 1 else get_month_str(mns[0]) + "至" + get_month_str(mns[-1])
        da = []
        for dept in set(list(ds.keys()) + list(de.keys()) + list(dt.keys())):
            d, e, t = ds.get(dept, 0), de.get(dept, 0), dt.get(dept, 0)
            if t > 0:
                a = (d + e) / 2
                r = round((t / a * 100), 2) if a > 0 else 0
                da.append({"一级组织": dept, "期初人数": d, "期末人数": e, "平均人数": round(a, 1), "离职人数": t, "离职率": r})
        da = sorted(da, key=lambda x: x["离职率"], reverse=True)
        td = []
        if len(all_t) > 0 and "离职类型" in all_t.columns:
            t = all_t.groupby("离职类型").size().reset_index(name="人数")
            t["占比"] = round(t["人数"] / tt * 100, 2) if tt > 0 else 0
            td = t.sort_values("人数", ascending=False).to_dict("records")
        rd = []
        if len(all_t) > 0 and "离职原因" in all_t.columns:
            r = all_t.groupby("离职原因").size().reset_index(name="人数")
            r["占比"] = round(r["人数"] / tt * 100, 2) if tt > 0 else 0
            rd = r.sort_values("人数", ascending=False).head(10).to_dict("records")
        ted = []
        if len(all_t) > 0 and "累计司龄（年）" in all_t.columns:
            tmp = all_t.copy()
            tmp["司龄段"] = tmp["累计司龄（年）"].apply(cat_tenure)
            t = tmp.groupby("司龄段").size().reset_index(name="人数")
            t["占比"] = round(t["人数"] / tt * 100, 2) if tt > 0 else 0
            ted = t.sort_values("人数", ascending=False).to_dict("records")
        ld = []
        if len(all_t) > 0 and "职级" in all_t.columns:
            tmp = all_t.copy()
            tmp["职级合并"] = tmp["职级"].apply(cat_level)
            l = tmp.groupby("职级合并").size().reset_index(name="人数")
            l["占比"] = round(l["人数"] / tt * 100, 2) if tt > 0 else 0
            ld = l.sort_values("人数", ascending=False).to_dict("records")
        return {"month": lbl, "avg": round(avg, 1), "start": ts, "end": te, "turn": tt, "rate": rate, "dept": da, "type": td, "reason": rd, "tenure": ted, "level": ld}

st.markdown('<div style="background: linear-gradient(135deg, #4F46E5 0%, #7C3AED 100%); padding: 20px; border-radius: 10px; margin-bottom: 20px;"><h2 style="color: white; margin: 0;">员工离职数据分析系统 V3.2</h2></div>', unsafe_allow_html=True)

st.header("数据导入")
f = st.file_uploader("上传Excel文件", type=["xlsx", "xls"])

if f:
    st.success("已上传: " + f.name)
    xlsx = pd.ExcelFile(f)
    p = Proc(xlsx)
    ms = p.months()
    
    if not ms:
        st.error("未找到期初数据，请检查文件格式！")
    else:
        options = []
        for m in ms: options.append((get_month_str(m), [m]))
        for i in range(len(ms)):
            for j in range(i + 1, len(ms)):
                if ms[j] - ms[i] == j - i:
                    options.append((get_month_str(ms[i]) + "至" + get_month_str(ms[j]), list(range(ms[i], ms[j] + 1))))
        
        sel = st.selectbox("选择月份", [o[0] for o in options])
        sel_mns = dict(options).get(sel, [1])
        
        if st.button("分析", type="primary"):
            r = p.analyze(sel_mns)
            
            st.subheader("关键指标")
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("分析月份", r["month"])
            c2.metric("平均人数", str(r["avg"]) + "人")
            c3.metric("离职人数", str(r["turn"]) + "人")
            c4.metric("离职率", str(r["rate"]) + "%")
            
            st.markdown("---")
            
            if r["dept"]:
                st.subheader("各部门离职分析")
                df_dept = pd.DataFrame(r["dept"])
                
                fig = go.Figure()
                fig.add_trace(go.Bar(x=df_dept["一级组织"], y=df_dept["离职人数"], name="离职人数", marker_color=COLORS["primary"], text=df_dept["离职人数"], textposition="outside"))
                fig.add_trace(go.Scatter(x=df_dept["一级组织"], y=df_dept["离职率"], name="离职率(%)", marker_color=COLORS["danger"], text=[f"{x:.2f}%" for x in df_dept["离职率"]], textposition="top center", yaxis="y2", mode="lines+markers+text", line=dict(width=3), marker=dict(size=10)))
                fig.update_layout(title=dict(text="各部门离职人数与离职率", font=dict(size=20)), xaxis=dict(title="部门", tickangle=-45, tickfont=dict(size=11), gridcolor="lightgray"), yaxis=dict(title="离职人数", title_font=dict(color=COLORS["primary"], size=14), tickfont=dict(size=12), gridcolor="lightgray"), yaxis2=dict(title="离职率(%)", title_font=dict(color=COLORS["danger"], size=14), overlaying="y", side="right", tickfont=dict(size=12)), legend=dict(x=0.5, y=1.12, xanchor="center", orientation="h", font=dict(size=13)), height=550, margin=dict(b=120), plot_bgcolor="white", paper_bgcolor="white")
                st.plotly_chart(fig, use_container_width=True, config=get_download_config(r["month"] + "_部门离职分析"))
                st.dataframe(df_dept, use_container_width=True, hide_index=True)
            
            st.markdown("---")
            
            if r["type"]:
                st.subheader("离职类型分布")
                df_type = pd.DataFrame(r["type"]).copy()
                for idx, row in df_type.iterrows():
                    if "主动" in str(row.get("离职类型", "")):
                        df_type.loc[idx, "离职类型"] = "主动离职"
                    elif "被动" in str(row.get("离职类型", "")):
                        df_type.loc[idx, "离职类型"] = "被动离职"
                merged = df_type.groupby("离职类型")["人数"].sum().reset_index()
                merged["占比"] = round(merged["人数"] / r["turn"] * 100, 2)
                
                fig_pie = px.pie(merged, values="人数", names="离职类型", title="离职类型占比", color_discrete_sequence=[COLORS["success"], COLORS["danger"]], hole=0.5)
                fig_pie.update_traces(textposition="outside", textinfo="label+percent+value", textfont=dict(size=14), marker=dict(line=dict(color="white", width=3)), hovertemplate="<b>%{label}</b><br>人数: %{value}人<br>占比: %{percent}<extra></extra>")
                fig_pie.update_layout(title=dict(text="离职类型分布", font=dict(size=20)), height=500, legend=dict(font=dict(size=14)))
                st.plotly_chart(fig_pie, use_container_width=True, config=get_download_config(r["month"] + "_离职类型分布"))
                st.dataframe(df_type, use_container_width=True, hide_index=True)
            
            st.markdown("---")
            
            if r["tenure"]:
                st.subheader("离职司龄分布")
                df_tenure = pd.DataFrame(r["tenure"])
                
                fig_tenure = px.pie(df_tenure, values="人数", names="司龄段", title="司龄分布", color_discrete_sequence=px.colors.qualitative.Set3, hole=0.5)
                fig_tenure.update_traces(textposition="outside", textinfo="label+percent+value", textfont=dict(size=13), marker=dict(line=dict(color="white", width=3)), hovertemplate="<b>%{label}</b><br>人数: %{value}人<br>占比: %{percent}<extra></extra>")
                fig_tenure.update_layout(title=dict(text="离职司龄分布", font=dict(size=20)), height=500, legend=dict(font=dict(size=14)))
                st.plotly_chart(fig_tenure, use_container_width=True, config=get_download_config(r["month"] + "_司龄分布"))
                st.dataframe(df_tenure, use_container_width=True, hide_index=True)
            
            st.markdown("---")
            
            if r["reason"]:
                st.subheader("离职原因分布")
                df_reason = pd.DataFrame(r["reason"])
                
                fig_reason = go.Figure()
                fig_reason.add_trace(go.Bar(x=df_reason["离职原因"], y=df_reason["人数"], marker_color=COLORS["warning"], text=[f"{d}人({p}%)" for d, p in zip(df_reason["人数"], df_reason["占比"])], textposition="outside", width=0.7))
                fig_reason.update_layout(title=dict(text="离职原因分布（纵向）", font=dict(size=20)), xaxis=dict(title="离职原因", tickangle=-30, tickfont=dict(size=11), gridcolor="lightgray"), yaxis=dict(title="人数", title_font=dict(size=14), tickfont=dict(size=12), gridcolor="lightgray"), height=max(450, len(df_reason) * 60), margin=dict(b=120, l=80), plot_bgcolor="white", paper_bgcolor="white")
                st.plotly_chart(fig_reason, use_container_width=True, config=get_download_config(r["month"] + "_离职原因分布"))
                st.dataframe(df_reason, use_container_width=True, hide_index=True)
            
            st.markdown("---")
            
            if r["level"]:
                st.subheader("职级分布（合并后）")
                df_level = pd.DataFrame(r["level"])
                
                fig_level = go.Figure()
                fig_level.add_trace(go.Bar(x=df_level["职级合并"], y=df_level["人数"], marker_color=COLORS["purple"], text=[f"{d}人({p}%)" for d, p in zip(df_level["人数"], df_level["占比"])], textposition="outside", width=0.7))
                fig_level.update_layout(title=dict(text="职级分布（人数+占比）", font=dict(size=20)), xaxis=dict(title="职级", tickfont=dict(size=14), gridcolor="lightgray"), yaxis=dict(title="人数", title_font=dict(size=14), tickfont=dict(size=12), gridcolor="lightgray"), height=450, plot_bgcolor="white", paper_bgcolor="white")
                st.plotly_chart(fig_level, use_container_width=True, config=get_download_config(r["month"] + "_职级分布"))
                st.dataframe(df_level.rename(columns={"职级合并": "职级"}), use_container_width=True, hide_index=True)
            
            st.markdown("---")
            
            # AI智能分析
            st.subheader("AI智能分析")
            with st.spinner("正在生成分析报告..."):
                analysis = generate_ai_analysis(r)
            st.markdown(analysis)
            
            # 导出选项
            col1, col2, col3 = st.columns(3)
            data_summary = get_data_summary(r)
            
            with col1:
                # PDF 导出
                try:
                    pdf_bytes = export_to_pdf(analysis, r["month"])
                    st.download_button(
                        label="📄 下载 PDF 报告",
                        data=pdf_bytes,
                        file_name=r["month"] + "离职分析报告.pdf",
                        mime="application/pdf"
                    )
                except Exception as e:
                    st.error(f"PDF导出失败: {str(e)}")
            
            with col2:
                # Word 导出
                try:
                    word_bytes = export_to_word(analysis, r["month"], data_summary)
                    st.download_button(
                        label="📝 下载 Word 报告",
                        data=word_bytes,
                        file_name=r["month"] + "离职分析报告.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Word导出失败: {str(e)}")
            
            with col3:
                # Markdown 导出
                md_content = f"# 员工离职分析报告\n\n{r['month']}\n\n{analysis}\n\n---\n\n## 原始数据\n\n{data_summary}"
                st.download_button(
                    label="📋 下载 Markdown",
                    data=md_content,
                    file_name=r["month"] + "离职分析报告.md",
                    mime="text/markdown"
                )
            
            st.markdown("---")
            
            # 导出Excel
            st.subheader("导出数据")
            out = BytesIO()
            with pd.ExcelWriter(out, engine="openpyxl") as w:
                pd.DataFrame([{"月份": r["month"], "期初人数": r["start"], "期末人数": r["end"], "平均人数": r["avg"], "离职人数": r["turn"], "离职率": str(r["rate"]) + "%"}]).to_excel(w, sheet_name="概览", index=False)
                if r["dept"]: pd.DataFrame(r["dept"]).to_excel(w, sheet_name="部门离职", index=False)
                if r["type"]: pd.DataFrame(r["type"]).to_excel(w, sheet_name="离职类型", index=False)
                if r["reason"]: pd.DataFrame(r["reason"]).to_excel(w, sheet_name="离职原因", index=False)
                if r["tenure"]: pd.DataFrame(r["tenure"]).to_excel(w, sheet_name="司龄分布", index=False)
                if r["level"]: pd.DataFrame(r["level"]).to_excel(w, sheet_name="职级分布", index=False)
            
            st.download_button(
                label="下载Excel报告",
                data=out.getvalue(),
                file_name=r["month"] + "离职分析.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
else:
    st.info("请上传Excel文件开始分析")
